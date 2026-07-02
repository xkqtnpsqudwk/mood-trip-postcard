from docx import Document
from pathlib import Path

PATH = Path(r"C:\epitech\middle-project\outputs\Mood Trip Postcard.docx")
doc = Document(PATH)

REPLACEMENTS = {
    "혼자 여행은 가고 싶지만 여행 계획을 세우는 과정은 부담스럽고": "혼자 여행은 가고 싶지만 어디로 가야 할지 고르는 과정은 부담스럽고",
    "계획 없이도, 느낌 있게 떠나는 AI 솔로 여행 큐레이터": "일정표 없이도, 느낌 있게 떠나는 AI 솔로 여행 큐레이터",
    "사용자는 여행 계획을 길게 세우지 않아도 된다.": "사용자는 복잡한 일정표를 길게 만들지 않아도 된다.",
    "계획은 AI에게 맡기고, 사용자는 오늘의 기분에 맞는 도시를 경험한다.": "일정표보다 오늘의 기분에 맞는 도시의 한 장면을 경험한다.",
    "이 서비스는 여행 전체 일정을 촘촘하게 설계하는 플래너가 아니다.": "이 서비스는 여행 전체 일정을 촘촘하게 설계하는 플래너가 아니라, 지금의 감정에 어울리는 장소를 골라주는 큐레이터이다.",
    "여행 계획을 세우기 귀찮다는 문제에서 출발한 서비스이므로": "여행지를 고르는 과정이 번거롭다는 문제에서 출발한 서비스이므로",
    "계획 짜는 건 귀찮고": "일정표 짜는 건 귀찮고",
    "계획 없이 나왔는데 오히려 더 편했다.": "정해둔 일정 없이 나왔는데 오히려 더 편했다.",
    "계획 없이 걸은 오후": "정해둔 길 없이 걸은 오후",
    "전체 여행 계획이나 예약 관리가 아니라": "전체 여행 일정 설계나 예약 관리가 아니라",
    "계획은 귀찮고": "일정표는 부담스럽고",
}


def replace_text(container):
    for paragraph in container.paragraphs:
        if not paragraph.runs:
            continue
        full_text = "".join(run.text for run in paragraph.runs)
        updated = full_text
        for old, new in REPLACEMENTS.items():
            updated = updated.replace(old, new)
        if updated != full_text:
            paragraph.runs[0].text = updated
            for run in paragraph.runs[1:]:
                run.text = ""


replace_text(doc)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            replace_text(cell)

doc.save(PATH)
print(PATH)
